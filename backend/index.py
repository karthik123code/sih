import os
from flask import Flask, request, jsonify, send_file, session
import instaloader
from docx import Document
# from flask_pymongo import PyMongo

app = Flask(__name__)

# MongoDB configuration
# app.config["MONGO_URI"] = "mongodb+srv://amrindersingh292004:YrUO6w88OMH3piB3@clusterinstascrapper.stjux.mongodb.net/"
# mongo = PyMongo(app)

@app.route('/', methods=['GET'])
def home():
    # Check if the user is logged in
    if 'username' in session:
        # Return a welcome message with the username
        return jsonify({'message': f'Welcome, {session["username"]}!'}), 200
    else:
        # Return a message indicating that the user is not logged in
        return jsonify({'message': 'You are not logged in'}), 401
    

@app.route('/login', methods=['POST'])
def login():
    # Get the request data
    data = request.get_json()

    # Extract the username and password from the request data
    username = data.get('username')
    password = data.get('password')

    # Check if the username and password match the expected values
    if username == 'admin' and password == 'Parse1234':
        # Return a success response with a JSON message
        return jsonify({'message': 'Login successful'}), 200
    else:
        # Return an error response with a JSON message
        return jsonify({'message': 'Invalid username or password'}), 401

@app.route('/instagram', methods=['POST'])
def insta_scraper():
    name = request.json.get('name')
    pwd = request.json.get('pwd')

    if not name or not pwd:
        return jsonify({'error': 'Username and password are required'}), 400

    try:
        # Initialize Instaloader
        L = instaloader.Instaloader()

        # Login to Instagram
        L.login(name, pwd)

        # Define the profile to scrape
        profile_name = name
        profile = instaloader.Profile.from_username(L.context, profile_name)

        # Create a new Document
        doc = Document()
        doc.add_heading(f'Instagram Profile Data: {profile_name}', 0)

        # Add profile information to the document
        doc.add_heading('Profile Information', level=1)
        doc.add_paragraph(f'Username: {profile.username}')
        doc.add_paragraph(f'Full Name: {profile.full_name}')
        doc.add_paragraph(f'Biography: {profile.biography}')
        doc.add_paragraph(f'Followers: {profile.followers}')
        doc.add_paragraph(f'Following: {profile.followees}')
        doc.add_paragraph(f'Number of Posts: {profile.mediacount}')

        # Add recent posts to the document
        doc.add_heading('Recent Posts', level=1)
        for post in profile.get_posts():
            doc.add_heading(post.date.strftime('%Y-%m-%d'), level=2)
            doc.add_paragraph(f"Post: {post.url}")
            doc.add_paragraph(f"Caption: {post.caption}")
            doc.add_paragraph(f'Likes: {post.likes}')
            doc.add_paragraph(f'Comments: {post.comments}')
            if post.video_url:
                doc.add_paragraph(f'Video URL: {post.video_url}')
            else:
                doc.add_paragraph(f'Image URL: {post.url}')
            doc.add_paragraph('')

        # Save the document
        doc.save('instagram_profile_data.docx')

        # Save the data to MongoDB
        # instagram_data = {
        #     'username': profile.username,
        #     'full_name': profile.full_name,
        #     'biography': profile.biography,
        #     'followers': profile.followers,
        #     'following': profile.followees,
        #     'mediacount': profile.mediacount,
        #     'posts': []
        # }
        # for post in profile.get_posts():
        #     instagram_data['posts'].append({
        #         'date': post.date,
        #         'url': post.url,
        #         'caption': post.caption,
        #         'likes': post.likes,
        #         'comments': post.comments,
        #         'video_url': post.video_url
        #     })
        # mongo.db.instagram_profiles.insert_one(instagram_data)

        return jsonify({'message': 'Data scraped and saved to instagram_profile_data.docx and MongoDB'})

    except instaloader.exceptions.BadCredentialsException:
        return jsonify({'error': 'Invalid username or password'}), 401
    except Exception as e:
        return jsonify({'error': f'An error occurred: {e}'}), 500

@app.route('/instagram/download', methods=['GET'])
def download_docx():
    # Check if the user is logged in
    if 'username' not in session:
        return jsonify({'error': 'You are not logged in'}), 401

    # Check if the DOCX file exists
    file_path = 'instagram_profile_data.docx'
    if not os.path.exists(file_path):
        return jsonify({'error': 'File not found'}), 404

    # Send the file as a download response
    return send_file(file_path, as_attachment=True, attachment_filename='instagram_profile_data.docx')
    
if __name__=="__main__":
    app.run(debug=True)